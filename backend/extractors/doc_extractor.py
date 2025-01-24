from docx import Document
from io import BytesIO
from .base_extractor import BaseExtractor

class DocExtractor(BaseExtractor):
    """Extracteur pour les fichiers Word (.doc, .docx)"""

    def extract(self, content: bytes) -> str:
        try:
            # Créer un objet BytesIO pour lire le contenu
            doc_file = BytesIO(content)

            # Charger le document Word
            doc = Document(doc_file)

            # Extraire le texte des paragraphes
            text_parts = []

            # Extraire le texte des paragraphes
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extraire le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            # Joindre tout le texte
            text = "\n".join(text_parts)
            print("Text extracted : \n", text)
            # print("======================================")
            # print("Text Cleaned", self.clean_text(text))
            # return self.clean_text(text)
            return text

        except Exception as e:
            raise Exception(f"Erreur lors de l'extraction du document Word: {str(e)}")